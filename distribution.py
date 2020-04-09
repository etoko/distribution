import pandas as pd
import matplotlib
import numpy as np

import matplotlib.pyplot as plt
from matplotlib import ticker as ptick
import time
import seaborn as sns

import docx

import datetime
import calendar
import os
from os.path import basename

import getopt, sys

import zipfile
from pathlib import Path
from shutil import copyfile
from shutil import copy

import tkinter as tk
from tkinter.filedialog import askopenfilename
from tkinter.filedialog import askopenfilenames
from tkinter.filedialog import asksaveasfilename
from tkinter.filedialog import askdirectory


"""
Script to generate docx reports from manifest export files during and after food/cash distribution.

"""

__version__ = '0.1.0'
__author__ = "Emmanuel Toko <emmanuel.toko@wfp.org>"

root = tk.Tk()
root.withdraw()

TABLE_STYLE = "Light Grid Accent 1"
CATEGORIES = {"f": "Food", "c": "Cash"}
DAILY_FOOD_RATION = {"Rice": 0.4, "CSB":0.05, "Pulses":0.08, "Salt":0.01, "Vegetable Oil":0.03}
#MONTHLY_ CASH_RATION = 31000
RECENT_DATE = datetime.date.today()
MONTHLY_CASH_RATION = 31000
RECENT_DATE_STRING = RECENT_DATE.strftime("%A, %d %B %Y")
df_gps = pd.DataFrame({
    "FDP": ["Youth Centre", "Magamaga"], 
    "Latitude":[1.93956692 , 1.92367864],
    "Longitude": [32.17859521, 32.16822177]
    })


def read_excel():
    manifest_file = askopenfilename(parent=root, defaultextension='.xlsx',\
        filetypes=[\
        ('Excel','*.xlsx'), \
        ('CSV','*.csv'), \
        ], title="Select Current Manifest")

    df = pd.read_excel(manifest_file)

    return df

def read_files(*args, **kwargs):
    """
    Read multiple manifests in excel format

    TODO: Consider adding the option to read csv format
    """
    default_title = "Select Files"
    infiles = askopenfilenames(parent=root, defaultextension='.xlsx',\
        filetypes=[\
        ('Excel','*.xlsx'), \
        ('CSV','*.csv'), \
        ], title=kwargs.get("title", default_title))


    data_frames = {}

    for infile in infiles:
        data_frames[infile] = pd.read_excel(infile)

    return data_frames


def create_document(d):
    doc = docx.Document()
    doc.add_heading(d["first_heading"], 0)

    return doc


def save(doc):
    file_name = asksaveasfilename()
    timestr = time.strftime("%Y%m%d-%H%M%S")
    #file_name = file_name+"-"+timestr+".csv"
    file_name = file_name+"-"+timestr+".docx"
    doc.save(file_name)
    file_name = '"' + file_name + '"'
    os.system(file_name)
    #print(file_name)


def export_excel(df):
    file_name = asksaveasfilename()
    timestr = time.strftime("%Y%m%d-%H%M%S")
    #file_name = file_name+"-"+timestr+".csv"
    file_name = file_name+"-"+timestr+".xlsx"
    df.to_excel(file_name, index=False)
    file_name = '"' + file_name + '"'
    os.system(file_name)
    #print(file_name)



def sanitise_table_header(df):
    cols = []
    for col in df.columns:
        #Watch out for strings in header and leave as is
        if type(col) == str:
            cols.append(col)
            continue
        elif type(col) == np.int64 or int:
            cols.append(str(col))
            continue
        else:
            cols.append(col.strftime("%d-%m-%Y"))

    df.columns = cols

    return df

def is_food(name):
    """
    Function checks the manifest name and establishes if it is a food manifest
    or cash manifest.

    returns True if a food manifest
    """
    name = name.split("-")
    manifest_type = CATEGORIES.get(name[3])
    if manifest_type == "Food":
        return True
    else:
        return False

    return False

def get_settlement(name):
    """
    Function checks the manifest name and retrieves the settlement name.

    returns a string containing the name of the settlement
    """
    if name is None:
        return None

    settlement = {}
    name = name.split("-")
    settlement["name"] = name[0].title() #return the settlement name
    benefit = name[3].strip().upper() #retrieve the benefit type
    benefit = "Food" if benefit == "F" else "Cash"
    settlement["modality"] = benefit
    settlement["fdp"] = name[4].split("_")[0].title() #retrieve name of the fdp
    settlement["cycle"] = name[1]

    return settlement


def concatenate(*args, **kwargs):
    """
    combine multiple manifest files from the same/different distribution 
    cyle/settlement/fdp into a single file.
    """
    data_frames = None

    if args:
        data_frames = args
    else:
        data_frames = read_files()

    frames = []

    for key, frame in enumerate(data_frames):
        file_name = basename(frame)
        settlement = get_settlement(file_name)
        frame = pd.read_excel(frame)
        frame["settlement"] = settlement["name"]
        frame["fdp"] = settlement["fdp"].split("_")[0].title()
        frame["modality"] = settlement["modality"]
        #frame["cycle"] = settlement["cycle"]
        frame.loc[:, "CreateDate"] = pd.to_datetime(frame.CreateDate)
        frame.loc[:,"cycle"] = frame.CreateDate.apply(lambda x: x.month)
        frames.append(frame)
        #print(frame)
    df = pd.concat(frames, sort=False)
    df.replace({'fdp': r'^You'}, {'fdp': 'Youth Centre'}, regex=True, inplace=True)
    sort_cols = ["CreateDate", "UpdateDate"]
    df.loc[:,'UpdateDate'] = pd.to_datetime(df.UpdateDate)
    #df.sort_values(by='', na_position="first", axis=1, inplace=True)

    #rename misplet FDP names
    df.replace(to_replace = "^Magamaga.*", value = "Magamaga", inplace = True, regex=True)
    df.replace(to_replace = "^Youth Centre.*", value = "Youth Centre", inplace = True, regex=True)

    if kwargs:
        for name, value in kwargs.items():
            if name == "export":
                if value:
                    export_excel(df) 
                else:
                    return df
            elif name == "frames":
                return df
            else:
                return df

    export_excel(df) 
    


def add_table_heading(df, doc, **kwargs):
    if kwargs:
        month = calendar.month_name[int(kwargs.get("mon"))]
        name = basename(kwargs.get("name"))
        name = name.split("-")
        try:
            heading = name[0] + " " + name[4].split("_")[0] + " " + CATEGORIES.get(name[3]) + " " + month
        except IndexError:
            heading = "Heading IndexErr"
        except TypeError:
            heading = "Heading (TypeError)"

        doc.add_heading(heading)
    else:
        doc.add_heading("Table Header Here")



def add_table(df, doc, **kwargs):
    doc_table = doc.add_table(df.shape[0]+1, df.shape[1])
    font_size = docx.shared.Pt(8)
    doc_table.style = TABLE_STYLE
    doc_table.style.font.size = font_size


    for i in range(df.shape[-1]):
        doc_table.cell(0,i).text = df.columns[i]
    
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            doc_table.cell(i+1,j).text = str(df.values[i,j])

    print("aDDED TABLE")


def add_plot(frames, doc):
    """
    Generate plot
    """
    dfs = pd.DataFrame()
    #Retrieve the totals of each frame
    for df in frames:
        frame = pivot_table(df)
        #total1 = frame.loc["Total"]
        #dfs.append(total1, ignore_index=True)
        #dfs.append(frame, ignore_index = True)
        frame = frame[-1:]
        print(frame.index)
        frame.plot.bar(x="Family Size")
        plt.savefig("chart.png")
        doc.add_picture("chart.png")

    print(dfs)
    #dfs.plot(kind="line")
    #plt.savefig("distribution_chart.png")    
    #doc.add_picture("distribution_chart.png")

def pivot_table(df, **kwargsl):
    df = df[df.Status=="Collected"]
    df.loc[:, "Family Size"] = df["ProcessingGroupSize"]
    df.loc[:, "UpdateDate"] = pd.to_datetime(df["UpdateDate"])
    df.loc[:, "UpdateDate"] = df.UpdateDate.dt.date
    #df["UpdateDate"] = df["UpdateDate"].dt.strftime("%d-%m-%Y")
    df = df.pivot_table(\
        index=["Family Size",], \
        values="ManifestGuid", columns=["UpdateDate",],\
        fill_value = 0,\
        aggfunc=len)

    print(df)
    df.loc["Total"] = df.sum()
    df.insert(loc=0, column="Family Size", value=df.index)
    
    cols = []
    for col in df.columns:
        #Watch out for strings in header and leave as is
        if type(col) == str:
            cols.append(col)
            continue
        cols.append(col.strftime("%d-%m-%Y"))

    df.columns = cols
    return df

def pivot_by_status(df, doc):
    #Pivot segregating beneficiaries by distribution status
    df.loc[:, "UpdateDate"] = pd.to_datetime(df["UpdateDate"])
    recent_date = df.UpdateDate.max().date()
    df = df[df["UpdateDate"].dt.date == recent_date]
    df_pivot = df.pivot_table(
            aggfunc=len, \
            index=["ProcessingGroupSize"], \
            values="ManifestGuid", \
            fill_value=0,\
            columns=["Status",])
    doc.add_heading("Households Served on " + recent_date.strftime("%A, %d %B %y"), level=3)
    df_pivot.loc["Total"] = df_pivot.sum()
    df_pivot.insert(0, "Family Size", df_pivot.index)
    add_table(df_pivot, doc)


def pt_food(df, doc, **kwargs):
    """
    Generate two pivot tables;
    first pivot is a break down of the beneficiaries by their status
    second pivot is the total benefit to beneficiaries with status "Collected"
    """
    #Distribution Table for 'Collected' beneficiaries
    df_distr = df[df.Status=="Collected"]
    df_distr = df_distr.pivot_table(\
            index=["ProcessingGroupSize",],\
            values="ReferenceID",
            aggfunc=len,
            fill_value = 0
            )
    df_distr.insert(loc=0, column="Family Size", value=df_distr.index)
    df_distr_cols = ["Family Size", "Collected"]
    df_distr.columns = df_distr_cols
    df_distr["Population"] = df_distr["Family Size"] * df_distr["Collected"]
    
    daily_ration = {"rice": 0.4, "CSB":0.05, "Pulses":0.08, "Salt":0.01, "Vegetable Oil":0.03}
    #monthly_ration = daily_ration * 30
    rations = pd.Series(daily_ration)
    df_distr["CSB"] = (df_distr.Collected * df_distr.index)*rations.CSB * 30
    df_distr["Pulses"] = (df_distr.Collected * df_distr.index)*rations.Pulses * 30
    df_distr["Salt"] = (df_distr.Collected * df_distr.index)*rations.Salt * 30
    df_distr["Vegetable_Oil"] = (df_distr.Collected * df_distr.index)*rations["Vegetable Oil"] * 30
    df_distr["Cereal"] = (df_distr.Collected * df_distr.index)*rations.rice * 30
    
    doc.add_heading("Population served", level=3)
    df_distr_copy = df_distr.copy()
    df_distr.loc["Total"] = df_distr.sum()
    df_distr.loc[:,"Total"] = df_distr.sum()
    #Format figures to include the thousandth separator
    df_distr = df_distr.round(2)
    df_distr.loc[:, "Cereal"] = df_distr["Cereal"].map("{:,}".format)
    df_distr.iloc[-1,0] = "Total"
    add_table(df_distr, doc)

    #Drop total row to exclude it from graph
    #df_distr.iloc[-1,:].drop(index=0,inplace=True)

    food = df_distr_copy[["Pulses", "Salt", "Vegetable_Oil", "CSB", "Cereal"]]
    food.index.name = "Family Size"
    food.plot(kind="bar", stacked=True, rot=0)
    ax = plt.gca()
    ax.set_ylabel("Kilograms")
    ax.grid(which='major', alpha=0.9)
    ax.grid(which='minor', alpha=0.4)
    ax.minorticks_on()
    plt.savefig("food_report.png")
    chart_heading = "Stacked chart showing quantities of food distributed to beneficiaries"
    doc.add_heading(chart_heading, level=3)
    doc.add_picture("food_report.png")

    if kwargs:
        df_unsorted = kwargs.get("unsorted_df")
        g = df_unsorted.pivot_table(index="ProcessingGroupSize", columns="Status",
            values="ManifestGuid", fill_value=0, aggfunc=len)
        g.index.name = "Family Size"
        total = df_unsorted.pivot_table(index="ProcessingGroupSize",
            values="ManifestGuid", aggfunc=len)
        g.loc[:, 'Total Households'] = total.ManifestGuid
        g.loc[:,"Population Served"] = g.index * g["Collected"]
        g['Percentage Served'] = (g["Collected"] / g["Total Households"])
        g["Percentage Served"] = g["Percentage Served"].map("{:.2%}".format)
        df_plot = g[["Collected", "Total Households"]]
        df_plot.columns = ("Collected", "Eligible")
        df_plot.plot(kind="bar", rot=0)
        ax = plt.gca()
        ax.set_ylabel("Households")
        ax.grid(which='major', alpha=0.9)
        ax.grid(which='minor', alpha=0.4)
        ax.minorticks_on()
        doc.add_heading("Cumulative Data", level=1)
        doc.add_heading("Total Households served from " + \
                kwargs.get("earliest_date"), level=3)
        g.insert(0, "Family Size", g.index)
        add_table(g, doc)
        plt.savefig("food_report.png", format="png")
        doc.add_heading("Clustered graph showing the percentage of \
                served vs eligible households", level=3)
        doc.add_picture("food_report.png")


def pt_cash(df, doc, *args, **kwargs):
    #Distribution Table
    df_distr = df[df.Status=="Collected"]
    df_distr = df_distr.pivot_table(\
            index=["ProcessingGroupSize",],\
            values="ReferenceID",
            aggfunc=len
            )
    df_distr_cols = ["Collected"]
    df_distr.index.name = "Family Size" 
    df_distr.columns = df_distr_cols
    #df_distr["Served Beneficiaries"] = df_distr["Family Size"] * df_distr["Collected"]
    df_distr["Served Beneficiaries"] = df_distr.index * df_distr["Collected"]
    
    daily_ration = {"rice": 0.4, "CSB":0.05, "Pulses":0.08, "Salt":0.01, "Vegetable Oil":0.03}
    #monthly_ration = daily_ration * 30
    rations = pd.Series(daily_ration)
    df_distr["Cash Distributed"] = (df_distr.Collected * df_distr.index)* MONTHLY_CASH_RATION
    df_distr.insert(loc=0, column="Family Size", value=df_distr.index)
    
    if kwargs:
        if kwargs.get("total"):
            return df_distr

    food = df_distr[["Cash Distributed"]]
    food.plot(kind="bar", stacked=True, rot=0)
    ax = plt.gca()
    ax.set_ylabel("Cash")
    ax.grid(which='major', alpha=0.9)
    ax.grid(which='minor', alpha=0.4)
    ax.yaxis.set_major_formatter(ptick.ScalarFormatter(useMathText=True)) 
    ax.minorticks_on()
    #plt.show()
    
    doc.add_heading("Distribution Status", level=1)
    df_distr_copy = df_distr.copy()
    #df_distr.loc['Total'] = df_distr_copy.sum()
    df_distr = df_distr.round(2)
    
    cash = df_distr["Cash Distributed"].sum()
    df_distr.loc["Total"] = df_distr.sum()
    df_distr.iloc[-1,0] = "Total"
    
    distr_text = "The total amount of cash distributed is as follows: {0:,}".format(cash)
    doc.add_paragraph(distr_text)
    df_distr["Cash Distributed"] = df_distr["Cash Distributed"].apply(lambda x: "{:,}".format(x))
    add_table(df_distr, doc)
    
    plt.savefig("stacked_food_chart.png")
    doc.add_picture("stacked_food_chart.png")
    try:
        doc.add_picture("food_report.png")
    except FileNotFoundError as err:
        pass

    df_all = df.copy()
    df_all = df_all.pivot_table(\
            index=["ProcessingGroupSize"],\
            values="ManifestGuid", \
            columns="Status", \
            aggfunc=len
            )
    df_all.index.name = "Family Size"
    df_all.insert(loc=0, column="Family Size", value=df_all.index)
    df_all.replace(np.nan, 0, inplace=True)
    df_all.loc["Total"] = df_all.sum()
    doc.add_heading("Distribution Summary", level=1)
    add_table(df_all, doc)
    

def litigation(df_litigation, doc, *args, **kwargs):
    """
    Process litigation dataframes
    """
    ref_row_index = np.where(df_litigation.isin(["Reference ID"]) == True)[0][0]
    df_litigation.drop(df_litigation.columns[8:df_litigation.shape[1]], \
            axis=1, inplace=True)
    off = [2,5,6]
    df_litigation.drop(df_litigation.columns[off], axis=1, inplace=True)
    cols = df_litigation.iloc[ref_row_index].tolist()
    cols = [str(title).strip() for title in cols]
    df_litigation.columns = cols
    #The Last row of this table contains the total so exclude this row
    df_litigation = df_litigation[ref_row_index+1:df_litigation.shape[0]-1]
    df_litigation.loc[:, "Date Time"] = pd.to_datetime(df_litigation["Date Time"])
    df_litigation.loc[:, "Date"] = df_litigation["Date Time"].dt.date
    #df_litigation.to_csv("lit.csv", index=False)

    #Segregate the months
    months = df_litigation["Date Time"].dt.month.value_counts()
    for month in months.index:
        df_lit_month = df_litigation[df_litigation["Date Time"].dt.month == month]
        df_lit_pivot = df_lit_month.pivot_table(columns="Date",\
            values= "Date Time", index=["Case", "New Status",], aggfunc=len)
        df_lit_pivot = df_lit_pivot.replace(np.nan, 0)

        df_lit_pivot = sanitise_table_header(df_lit_pivot)
        litigation_table_heading = "Litigation"
        if kwargs:
            file_name = basename(kwargs.get("name"))
            file_name = file_name.split("-")
            litigation_table_heading = file_name
            try:
                litigation_table_heading = "Litigation " + \
                    calendar.month_name[int(month)] + \
                    " " + file_name[1]+ " " + \
                    file_name[5].split("-")[0] + " " + CATEGORIES.get(file_name[4])
            except IndexError:
                litigation_table_heading = file_name

        doc.add_heading(litigation_table_heading)
        df_lit_pivot.reset_index(level=0, inplace=True)
        add_table(df_lit_pivot, doc)


def verification(df, doc):
    if df.empty:
        doc.add_heading("No Uncollected Beneficiaies Exist", level=2)
        return

    df.loc[:,"UpdateDate"] = pd.to_datetime(df["UpdateDate"])
    df_verification = df[df["Status"] == "Verified"]
    if df_verification.empty:
        doc.add_heading("No Uncollected Beneficiaies Exist", level=2)
        return

    verification_columns = ["ReferenceID", "ProcessingGroupNumber", \
        "ProcessingGroupSize", "VerifiedBy", "UpdateDate"]
    df_verification = df_verification[verification_columns]
    df_ver_day = df_verification.copy()
    df_ver_day.loc[:,"Date"] = df_ver_day["UpdateDate"].dt.date
    df_ver_day_pv = df_ver_day.pivot_table(aggfunc=len, columns="Date", \
            values="ProcessingGroupNumber", index="ProcessingGroupSize",\
            fill_value=0)
    df_ver_day_pv = sanitise_table_header(df_ver_day_pv)
    df_ver_day_pv.loc["Total"] = df_ver_day_pv.sum()
    df_ver_day_pv.insert(0, "Family Size", df_ver_day_pv.index)
    add_table(df_ver_day_pv, doc)
    doc.add_heading("Total number of un-exited beneficiaries",level=3)
    df_verification.sort_values("UpdateDate", ascending=False, inplace=True)
    df_verification.insert(0, "#", range(1, df_verification.shape[0]+1))
    add_table(df_verification, doc)


def daily(): 
    """
    Generate daily report
    """
    d = {"first_heading":"Daily Distribution Report"}
    doc = create_document(d)
    
    #Distribution Reports
    #frames = read_files()

    #Store for dataframes sorted according to month
    sorted_frames = {0:[], 1:[]}

    #num_manifests = len(frames)

    def create_pivot():
        df = concatenate(export=False)

        df.loc[:, "UpdateDate2"] = pd.to_datetime(df.UpdateDate.dt.date)
        pivot = df.pivot_table(values=["ManifestGuid", "ProcessingGroupSize", "UpdateDate2"], 
                index=["modality", "settlement"], 
                aggfunc={"ManifestGuid": len, "UpdateDate2": [np.min, np.max], "ProcessingGroupSize": np.sum})
        collected_df = df[df.Status == "Collected"]
        collected_df.loc[:, "UpdateDate2"] = pd.to_datetime(collected_df.UpdateDate.dt.date)
        collected_pivot = collected_df.pivot_table(values=["ManifestGuid", "ProcessingGroupSize", 
            "UpdateDate2"], index=["modality", "settlement"], 
            aggfunc={"ManifestGuid": len, "UpdateDate2": [np.min, np.max], "ProcessingGroupSize": np.sum})
        merged_pivot = pivot.merge(collected_pivot, left_index=True, right_index=True)
        merged_pivot = merged_pivot.reset_index(level=[0,1])
        merged_pivot['index_heading'] = merged_pivot.index
        merged_pivot = merged_pivot.drop(merged_pivot.columns[[-3, -2,-1]], axis=1)
        cols = ["Modality", "Settlement", "Planned Households", 
                "Planned Population", "End Date", "Start Date", 
                "Served Households", "Served Population"]
        merged_pivot.columns = cols
        cols = ["Modality", "Settlement", "Planned Households", 
                "Planned Population", "Served Households", "Served Population", 
                "Start Date", "End Date"]
        merged_pivot = merged_pivot[cols]
        merged_pivot.loc[:, "Planned Households"] = merged_pivot["Planned Households"].map("{:,}".format)
        merged_pivot.loc[:, "Planned Population"] = merged_pivot["Planned Population"].map("{:,}".format)
        merged_pivot.loc[:, "Served Population"] = merged_pivot["Served Population"].map("{:,}".format)
        merged_pivot.loc[:, "Served Households"] = merged_pivot["Served Households"].map("{:,}".format)
        merged_pivot.loc[:, "Start Date"] = pd.to_datetime(merged_pivot["Start Date"])
        merged_pivot.loc[:, "Start Date"] = merged_pivot["Start Date"].dt.strftime("%d-%b-%Y")
        merged_pivot.loc[:, "End Date"] = pd.to_datetime(merged_pivot["End Date"])
        merged_pivot.loc[:, "End Date"] = merged_pivot["End Date"].dt.strftime("%d-%b-%Y")
        merged_pivot = sanitise_table_header(merged_pivot)
        doc.add_heading("Executive Summary", level=1)
        add_table(merged_pivot, doc)


        #Detail by status
        df_pivot = df.pivot_table(aggfunc={len}, index=["modality", "settlement", "fdp"], 
                values=["ManifestGuid"], columns=["Status"], fill_value=0)
        #df_pivot.loc[:, "Total"] = df_pivot.sum()
        df_pivot = df_pivot.reset_index(level=[0, 1]) 
        df_pivot.loc[:, "FDP"] = df_pivot.index

        #Format the table columns to be more readible
        new_cols = []
        cols = df_pivot.columns
        for index, col in enumerate(cols):
            if index == 0 or index == 1 or index == (len(cols)-1):
                col_name = col[0].upper()
                new_cols.insert(index, col_name)
            else:
                col_name = col[-1].upper()
                new_cols.insert(index, col_name)

        df_pivot.columns = new_cols
        #Move the column for FDP to the third position
        type(new_cols)
        df_pivot.columns = new_cols.insert(2, new_cols.pop(-1))
        df_pivot = sanitise_table_header(df_pivot)
        doc.add_heading("Distribution by Status", level=1)
        add_table(df_pivot, doc)

        save(doc)
        #settlements = df.settlements.unique()
        #for settlement in settlements:
        #    settlement_dist = frame[frame.settlement == settlement]
        #    start_date = frame.UpdateDate.min.date()
        #    end_date = frame.UpdateDate.max.date()

        
    create_pivot()
    
    #if num_manifests > 1:
    #    create_pivot()
    #    return
    #else:
    #        df = concatenate(frames=True)

    return


    for key, frame in frames.items():
        #TODO: Add distribution chart for each distribution or merged chart for all distributions
        frame["UpdateDate"] = pd.to_datetime(frame["UpdateDate"])
        months = frame.UpdateDate.dt.month.value_counts()
        for month in months.index:
            RECENT_DATE = frame.UpdateDate.max().date()
            EARLIEST_DATE = frame.UpdateDate.min().date()
            RECENT_DATE_STRING = RECENT_DATE.strftime("%A, %d %B %Y")
            EARLIEST_DATE_STRING = EARLIEST_DATE.strftime("%A, %d %B %Y")
            df_month = frame[frame.UpdateDate.dt.month == month]; 
            df_recent_date = df_month[df_month.UpdateDate.dt.date == RECENT_DATE]
            
            if df_recent_date.empty:
                #TODO Explore better handling of missing data
                continue

            df_mon = pivot_table(df_recent_date)
            doc.add_heading(RECENT_DATE_STRING, level=1)

            if is_food(key):
                pivot_by_status(df_recent_date, doc)
                pt_food(df_recent_date, doc, unsorted_df=frame, 
                        earliest_date=EARLIEST_DATE_STRING)
            else:
                pivot_by_status(df_recent_date, doc)
                pt_cash(df_recent_date, doc)
            #sorted_frames.get(month, month).append(df_month);
            #print(sorted_frames)
            df_month_pv = pivot_table(df_month)
            #add_table_heading(df_month_pv, doc, mon=month, name=key)
            doc.add_heading("Households served per day from " + \
                    EARLIEST_DATE_STRING + " to " + RECENT_DATE_STRING
                    , level=3)
            add_table(df_month_pv, doc, mon=month, name=key)
            #doc.add_heading("Verified But not Collected", level=1)
            doc.add_heading("Beneficiaries not exited", level=3)
            verification(frame, doc)


    #add_plot(frames.items(), doc)
    dfs_litigation = read_files(title="Select Litigation Files")
    for key, df_litigation in dfs_litigation.items():
        litigation(df_litigation, doc, name=key)

    save(doc)

def weekly():
    print("Weekly")

def monthly():
    """
    Generate report for the month. The month/s is/are determined based on the 
    UpdateDate column of the manifest export sheet.
    """
    d = {"first_heading":"Monthly Distribution Report"}
    doc = create_document(d)
    
    #Distribution Reports
    frames = read_files()

    #Store for dataframes sorted according to month
    sorted_frames = {0:[], 1:[]}

    for key, frame in frames.items():
        #TODO: Add distribution chart for each distribution or merged chart for all distributions
        frame["UpdateDate"] = pd.to_datetime(frame["UpdateDate"])
        months = frame.UpdateDate.dt.month.value_counts()
        for month in months.index:
            doc.add_heading(calendar.month_name[int(month)], level=1)
            df_month = frame[frame.UpdateDate.dt.month == month]; 
            #sorted_frames.get(month, month).append(df_month);
            #print(sorted_frames)
            df_mon = pivot_table(df_month)
            add_table_heading(df_mon, doc, mon=month, name=key)
            add_table(df_mon, doc, mon=month, name=key)
            if is_food(key):
                pt_food(df_month, doc)
            else:
                pt_cash(df_month, doc)
        #add_table(df, doc)
    
    #Read Litigation Files
    #add_plot(frames.items(), doc)
    dfs_litigation = read_files(title="Select Litigation Files")
    for key, df_litigation in dfs_litigation.items():
        litigation(df_litigation, doc, name=key)

    save(doc)


def cycle():
    """
    Generate statistics based on the distribution cycles rather than month
    """
    cycle = input("Please enter the distribution cycle: ")
    cycle = int(cycle)
    print("Processing manifests ...")
    frames = concatenate(export=False)
    print("...[Complete]")
    d = {"first_heading": "Report for Distribution Cycle {0:d}".format(cycle)}
    doc = create_document(d)
    frames.loc[:, "cycle"] = pd.to_numeric(frames.cycle)
    d = {"first_heading": calendar.month_name[cycle] + " Distribution Report"}
    doc.add_heading("Distribution Trends", level=2)
    df_cash_collected = frames[frames.Status == "Collected"]
    df_cash_collected = df_cash_collected[df_cash_collected.benefit == "Cash"]
    df_cash_collected.loc[:,"Cash"] = df_cash_collected.loc[:,"ProcessingGroupSize"] * MONTHLY_CASH_RATION
    CEREAL_RATION = 0.4*31
    PULSE_RATION = 31 * 0.08
    VEGETABLE_OIL_RATION = 31*0.03
    SALT_RATION = 31*0.01
    CSB_RATION = 31*0.05
    df_cash_collected_pv = df_cash_collected.pivot_table(columns="cycle", aggfunc=np.sum, values = "Cash") 
    df_cash_collected_pv = sanitise_table_header(df_cash_collected_pv)
    add_table(df_cash_collected_pv, doc)
    df_food_collected = frames[frames.benefit=="Food"]
    #df_food_collected = frames[""]


    frames_hh_benefit_pivot = frames.pivot_table(aggfunc=[len, np.sum], columns="cycle", index="benefit", values="ProcessingGroupSize")
    frames_popn_benefit_pivot = frames.pivot_table(aggfunc=[len, np.sum], columns="cycle", index=["benefit", "Status"], values="ProcessingGroupSize")
    #frames_pivot = frames.pivot_table(index, columns="cycle", aggfunc=len)
    frames_hh_benefit_pivot.insert(0, "Value", frames_hh_benefit_pivot.index)
    doc.add_heading("Number of Households and Population Receiving Assistance by month", level=3)

    grand_total = pd.concat(
            [frames_hh_benefit_pivot, frames_popn_benefit_pivot]
            ).append(
                    frames_hh_benefit_pivot.sum().rename(('Grand', 'Total'))
                    )
    grand_total = sanitise_table_header(grand_total)
    add_table(grand_total, doc)

    frames_hh_benefit_pivot = sanitise_table_header(frames_hh_benefit_pivot)

    frames_popn_benefit_pivot.insert(0, "Value", frames_popn_benefit_pivot.index)
    frames_popn_benefit_pivot = sanitise_table_header(frames_popn_benefit_pivot)
    #add_table(frames_hh_benefit_pivot, doc)
    #add_table(frames_popn_benefit_pivot, doc)
    df = frames[frames.cycle == cycle]
    df_pivot = pivot_table(df)
    add_table_heading(df_pivot, doc, mon=cycle, name="") #TODO Add sensible name
    add_table(df_pivot, doc, mon=cycle, name="")

    df_food = df[df.benefit == "Food"]
    df_cash = df[df.benefit == "Cash"]
    
    #df_cash_pv = pt_cash(df_cash, doc, total = True)
    
    df_food_collected = df_food[df_food.Status == "Collected"]
    df_food_collected_pivot = df_food_collected.pivot_table(aggfunc=np.sum, index="ProcessingGroupSize", \
            values="Notes", columns="cycle")
    df_food_collected_pivot = sanitise_table_header(df_food_collected_pivot)
    #add_table(df_cash_pv, doc)
    df_cash_collected = df_cash[df_cash.Status == "Collected"]

    if not df_food.empty:
        pt_food(df_food, doc)
    if not df_cash.empty:
        pt_cash(df_cash, doc)

    save(doc)


def search_manifest():

    #manifest_file = askopenfilename(parent=root, defaultextension='.zip',\
    #    filetypes=[\
    #    ('Zip','*.zip'), \
    #    ], title="Select Zip File")
    #print(manifest_file)
    #z = zipfile.ZipFile(manifest_file)
    #for f in z.namelist():
    #    dirname = os.path.splitext(f)[0]
    #    print(dirname)
    src_directory = askdirectory()
    file_list = []
    
    for path in Path(src_directory).rglob('[a-z]*.xlsx'):
        file_list.append(path)

    dst_directory = askdirectory()

    for src in file_list:
        copy(src, dst_directory)


def main():
    try:
        opts, args = getopt.getopt(sys.argv[1:], "ho:v:c", ["help", "period=", "concat", "search", "s"])
    except getopt.GetoptError as err:
        # print help information and exit:
        print(err)  # will print something like "option -a not recognized"
        print("Implement Output here")
        sys.exit(2)
    output = None
    verbose = False
    for o, a in opts:
        if o == "-v":
            verbose = True
        elif o in ("-h", "--help"):
            print("HELP")
            sys.exit()
        elif o in ("-o", "--period"):
            if a.lower() == "weekly":
                weekly()
            elif a.lower() == "monthly":
                monthly()
            elif a.lower() == "daily":
                daily()
            elif a.lower() == "cycle":
                cycle()
        elif o in ("-s", "--search"):
            search_manifest()
        elif o in ("--concat", "-c"):
            concatenate()
        else:
            daily()
    # ...

if __name__ == "__main__":
    main()
